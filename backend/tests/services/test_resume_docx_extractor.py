import tempfile
from pathlib import Path
from unittest import TestCase
from unittest.mock import patch

from docx import Document
from docx.oxml import parse_xml

from app.services.resume_docx_extractor import (
    ResumeDocxExtractor,
    ResumeDocxExtractorError,
)


class ResumeDocxExtractorTest(TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp_dir.cleanup)
        self.storage_root = Path(self.temp_dir.name)
        self.resume_dir = self.storage_root / "v2" / "resumes" / "2026" / "08"
        self.resume_dir.mkdir(parents=True)
        self.extractor = ResumeDocxExtractor()

    def save_document(self, document: Document, name: str = "resume.docx") -> Path:
        path = self.resume_dir / name
        document.save(path)
        return path

    def extract(self, path: Path) -> str:
        return self.extractor.extract(
            self.storage_root,
            path.relative_to(self.storage_root).as_posix(),
            path.stat().st_size,
        )

    @staticmethod
    def add_text_box(document: Document, text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph._p.append(
            parse_xml(
                f"""
                <w:r
                    xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                    xmlns:v="urn:schemas-microsoft-com:vml"
                >
                    <w:pict>
                        <v:shape id="ResumeTextBox">
                            <v:textbox>
                                <w:txbxContent>
                                    <w:p><w:r><w:t>{text}</w:t></w:r></w:p>
                                </w:txbxContent>
                            </v:textbox>
                        </v:shape>
                    </w:pict>
                </w:r>
                """
            )
        )

    def test_extracts_paragraphs_and_tables_in_document_order(self) -> None:
        document = Document()
        document.add_paragraph("个人简介")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "学校"
        table.cell(0, 1).text = "时间"
        table.cell(1, 0).text = "西安建筑科技大学"
        table.cell(1, 1).text = "2024-至今"
        document.add_paragraph("项目经历")
        path = self.save_document(document)

        text = self.extract(path)

        self.assertEqual(
            text,
            "个人简介\n学校\t时间\n西安建筑科技大学\t2024-至今\n项目经历",
        )

    def test_does_not_duplicate_merged_table_cells(self) -> None:
        document = Document()
        table = document.add_table(rows=1, cols=2)
        merged = table.cell(0, 0).merge(table.cell(0, 1))
        merged.text = "合并标题"
        path = self.save_document(document)

        self.assertEqual(self.extract(path), "合并标题")

    def test_extracts_text_from_word_text_boxes(self) -> None:
        document = Document()
        self.add_text_box(document, "文本框中的项目经历")
        path = self.save_document(document)

        self.assertEqual(self.extract(path), "文本框中的项目经历")

    def test_uses_standard_text_when_mammoth_fails(self) -> None:
        document = Document()
        document.add_paragraph("普通段落仍然可用")
        self.add_text_box(document, "文本框内容")
        path = self.save_document(document)

        with patch(
            "app.services.resume_docx_extractor.mammoth.extract_raw_text",
            side_effect=RuntimeError("simulated failure"),
        ):
            self.assertEqual(self.extract(path), "普通段落仍然可用")

    def test_rejects_document_without_effective_text(self) -> None:
        path = self.save_document(Document())

        with self.assertRaisesRegex(ResumeDocxExtractorError, "没有可提取"):
            self.extract(path)

    def test_rejects_damaged_docx(self) -> None:
        path = self.resume_dir / "resume.docx"
        path.write_bytes(b"not-a-docx")

        with self.assertRaisesRegex(ResumeDocxExtractorError, "损坏或无法读取"):
            self.extract(path)

    def test_rejects_archive_with_too_many_members(self) -> None:
        document = Document()
        document.add_paragraph("有效内容")
        path = self.save_document(document)

        with patch(
            "app.services.resume_docx_extractor.MAX_DOCX_ARCHIVE_ENTRIES",
            1,
        ):
            with self.assertRaisesRegex(ResumeDocxExtractorError, "文件结构过于复杂"):
                self.extract(path)

    def test_rejects_archive_with_excessive_uncompressed_size(self) -> None:
        document = Document()
        document.add_paragraph("有效内容")
        path = self.save_document(document)

        with patch(
            "app.services.resume_docx_extractor.MAX_DOCX_UNCOMPRESSED_BYTES",
            100,
        ):
            with self.assertRaisesRegex(ResumeDocxExtractorError, "解压后内容超过"):
                self.extract(path)

    def test_rejects_extracted_text_over_limit(self) -> None:
        document = Document()
        document.add_paragraph("超过限制的文字")
        path = self.save_document(document)

        with patch(
            "app.services.resume_docx_extractor.MAX_EXTRACTED_CHARACTERS",
            4,
        ):
            with self.assertRaisesRegex(ResumeDocxExtractorError, "文本超过安全长度"):
                self.extract(path)

    def test_rejects_size_mismatch(self) -> None:
        document = Document()
        document.add_paragraph("有效内容")
        path = self.save_document(document)

        with self.assertRaisesRegex(ResumeDocxExtractorError, "文件大小与上传记录不一致"):
            self.extractor.extract(
                self.storage_root,
                path.relative_to(self.storage_root).as_posix(),
                path.stat().st_size + 1,
            )

    def test_rejects_path_outside_private_namespace(self) -> None:
        outside = self.storage_root / "outside.docx"
        document = Document()
        document.add_paragraph("有效内容")
        document.save(outside)

        with self.assertRaisesRegex(ResumeDocxExtractorError, "文件路径无效"):
            self.extractor.extract(
                self.storage_root,
                "outside.docx",
                outside.stat().st_size,
            )
